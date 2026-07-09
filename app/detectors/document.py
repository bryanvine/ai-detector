"""Document handling: extract text (+ metadata context) from PDF/DOCX/TXT/MD.

The verdict for documents comes from the text pipeline run on the extracted
body. Producer/creator metadata is surfaced as context only (weight 0) — a
PDF exported by LaTeX or headless Chrome says something about the *export*
tool, not about who wrote the words.
"""
from __future__ import annotations

import io

from ..ensemble import Signal


class ExtractionError(Exception):
    pass


def _from_pdf(data: bytes) -> tuple[str, dict]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception as exc:
            raise ExtractionError("PDF is password-protected") from exc
    text = "\n\n".join((page.extract_text() or "") for page in reader.pages)
    meta = reader.metadata or {}
    info = {
        "pages": len(reader.pages),
        "producer": str(meta.get("/Producer", "") or ""),
        "creator": str(meta.get("/Creator", "") or ""),
    }
    return text, info


def _from_docx(data: bytes) -> tuple[str, dict]:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    props = document.core_properties
    info = {
        "author": props.author or "",
        "last_modified_by": props.last_modified_by or "",
        "revision": props.revision,
    }
    return "\n".join(parts), info


def _from_text(data: bytes) -> tuple[str, dict]:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc), {"encoding": enc}
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise ExtractionError("Could not decode file as text")


def extract(filename: str, data: bytes) -> tuple[str, Signal]:
    """Return (extracted_text, info_signal). Raises ExtractionError."""
    lower = (filename or "").lower()
    if lower.endswith(".pdf") or data[:5] == b"%PDF-":
        text, info = _from_pdf(data)
        kind = "PDF"
    elif lower.endswith((".docx", ".docm")):
        text, info = _from_docx(data)
        kind = "DOCX"
    elif data[:4] == b"PK\x03\x04":
        raise ExtractionError(f"Unsupported document type: {filename}")
    else:
        text, info = _from_text(data)
        kind = "text"

    text = text.strip()
    if len(text) < 120:
        raise ExtractionError(
            f"Only {len(text)} characters of text could be extracted — "
            "not enough to analyze (scanned/image PDFs are not OCRed)."
        )
    detail = f"{kind}, {len(text):,} chars extracted"
    extras = ", ".join(f"{k}={v}" for k, v in info.items() if v)
    if extras:
        detail += f" ({extras})"
    return text, Signal("doc_meta", "Document metadata", None, 0, detail, dict(info))
